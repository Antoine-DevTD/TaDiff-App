from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "doc"
OUTPUT.mkdir(parents=True, exist_ok=True)
TARGET = OUTPUT / "TaDiff_informations_juridiques_et_commerciales_a_completer.docx"

BLUE = "2455DC"
NAVY = "0B1427"
LIGHT_BLUE = "E8EFFF"
PALE = "F5F7FB"
LINE = "C9D4E6"
MUTED = RGBColor(86, 101, 128)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    tr_pr.append(table_header)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_heading(document, number, title, introduction=None):
    paragraph = document.add_paragraph()
    keep_with_next(paragraph)
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(f"{number}. {title}")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    if introduction:
        intro = document.add_paragraph(introduction)
        intro.paragraph_format.space_after = Pt(8)
        intro.paragraph_format.keep_with_next = True
        intro.runs[0].font.color.rgb = MUTED
        intro.runs[0].font.size = Pt(9.5)


def add_field_table(document, rows, widths=(6.3, 10.2)):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    header.cells[0].text = "Information demandée"
    header.cells[1].text = "Réponse de Tony"
    set_repeat_table_header(header)
    for index, cell in enumerate(header.cells):
        cell.width = Cm(widths[index])
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
    for label, suggestion in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(widths[0])
        cells[1].width = Cm(widths[1])
        cells[0].text = label
        cells[1].text = suggestion or "À compléter :"
        for cell in cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)
            cell.paragraphs[0].paragraph_format.space_before = Pt(1)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8.8)
        set_cell_shading(cells[0], PALE)
        if suggestion:
            set_cell_shading(cells[1], LIGHT_BLUE)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_question(document, question, suggestion=None, lines=1):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_border(cell)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    label = paragraph.add_run(question)
    label.bold = True
    label.font.size = Pt(9.3)
    if suggestion:
        proposal = cell.add_paragraph(f"Proposition actuelle à confirmer : {suggestion}")
        proposal.paragraph_format.space_after = Pt(3)
        proposal.runs[0].italic = True
        proposal.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
        proposal.runs[0].font.size = Pt(8.7)
    answer = cell.add_paragraph("Réponse : ")
    answer.paragraph_format.space_after = Pt(0)
    answer.runs[0].font.size = Pt(8.8)
    if lines > 1:
        answer.add_run("\n" * (lines - 1))
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_break(document):
    document.add_page_break()


document = Document()
section = document.sections[0]
section.top_margin = Cm(1.7)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

styles = document.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
styles["Normal"].font.size = Pt(9.5)
styles["Normal"].font.color.rgb = RGBColor.from_string(NAVY)
styles["Normal"].paragraph_format.space_after = Pt(5)

header = section.header
header_paragraph = header.paragraphs[0]
header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header_run = header_paragraph.add_run("TaDiff  |  Document de préparation")
header_run.bold = True
header_run.font.size = Pt(8.5)
header_run.font.color.rgb = RGBColor.from_string(BLUE)

footer = section.footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_paragraph.add_run("Informations confidentielles - à compléter avant validation juridique")
footer_run.font.size = Pt(7.5)
footer_run.font.color.rgb = MUTED

brand = document.add_paragraph()
brand.paragraph_format.space_after = Pt(28)
brand_run = brand.add_run("TADIFF")
brand_run.bold = True
brand_run.font.size = Pt(12)
brand_run.font.color.rgb = RGBColor.from_string(BLUE)

title = document.add_paragraph()
title.paragraph_format.space_after = Pt(10)
title_run = title.add_run("Informations juridiques et commerciales à compléter")
title_run.bold = True
title_run.font.size = Pt(25)
title_run.font.color.rgb = RGBColor.from_string(NAVY)

subtitle = document.add_paragraph(
    "Ce document rassemble les informations nécessaires pour finaliser les mentions légales, "
    "les conditions d'utilisation, les conditions de vente et la documentation RGPD de TaDiff."
)
subtitle.paragraph_format.space_after = Pt(18)
subtitle.runs[0].font.size = Pt(11)
subtitle.runs[0].font.color.rgb = MUTED

notice = document.add_table(rows=1, cols=1)
notice.alignment = WD_TABLE_ALIGNMENT.CENTER
notice_cell = notice.cell(0, 0)
set_cell_shading(notice_cell, LIGHT_BLUE)
set_cell_border(notice_cell, BLUE)
notice_cell.text = (
    "Mode d'emploi : compléter les champs bleus, cocher les options retenues et joindre les justificatifs "
    "indiqués à la fin. En cas de doute, écrire \"à confirmer avec le comptable ou le juriste\"."
)
for run in notice_cell.paragraphs[0].runs:
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(NAVY)

document.add_paragraph()
meta = document.add_table(rows=3, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
meta.autofit = False
for row, label in zip(meta.rows, ["Document complété par", "Date", "Version"]):
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(11.5)
    row.cells[0].text = label
    row.cells[1].text = "À compléter :"
    set_cell_shading(row.cells[0], PALE)
    for cell in row.cells:
        set_cell_border(cell)

add_heading(document, "1", "Identité de l'exploitant actuel", "Ces informations doivent correspondre exactement aux documents officiels d'ARKENCIEL.")
add_field_table(document, [
    ("Dénomination exacte", "ARKENCIEL Compagnie, orthographe à confirmer"),
    ("Forme juridique", "Association / société / autre :"),
    ("Adresse complète du siège", None),
    ("Numéro SIREN", None),
    ("Numéro SIRET", None),
    ("Numéro RNA ou RCS", None),
    ("Capital social, si applicable", None),
    ("Numéro de TVA ou mention d'exonération", None),
    ("Nom du représentant légal", None),
    ("Nom du directeur de publication", None),
    ("Téléphone professionnel", None),
    ("Adresse email juridique", "contact@tadiff.com à confirmer"),
    ("Adresse email RGPD", "contact@tadiff.com à confirmer"),
    ("Adresse email support", "contact@tadiff.com à confirmer"),
    ("Adresse email facturation", "contact@tadiff.com à confirmer"),
])

add_page_break(document)
add_heading(document, "2", "Offre bêta et facturation", "Les réponses de cette partie seront reprises dans les conditions générales de vente.")
add_question(document, "Quel organisme encaissera les premiers abonnements ?", "ARKENCIEL Compagnie")
add_question(document, "Le prix de 19,99 EUR est-il HT ou TTC ?", "[ ] HT    [ ] TTC")
add_question(document, "ARKENCIEL facture-t-elle la TVA ?", "[ ] Oui    [ ] Non    [ ] À confirmer avec le comptable")
add_question(document, "À qui l'offre est-elle destinée ?", "[ ] Professionnels et associations uniquement    [ ] Particuliers également")
add_question(document, "Quelle est la durée de l'abonnement ?", "Mensuel, sans engagement minimum")
add_question(document, "L'abonnement se renouvelle-t-il automatiquement chaque mois ?", "[ ] Oui    [ ] Non")
add_question(document, "Quand la résiliation prend-elle effet ?", "À la fin de la période déjà payée")
add_question(document, "Une période commencée peut-elle être remboursée ?", "Non, sauf double paiement, obligation légale ou manquement de TaDiff")
add_question(document, "Le tarif bêta sera-t-il conservé après la bêta ?", "Préciser la durée et les conditions", lines=2)
add_question(document, "Quel délai de préavis appliquer avant une hausse de prix ?", "30 jours")

add_page_break(document)
add_heading(document, "3", "Limites comprises dans l'offre")
add_field_table(document, [
    ("Nombre d'utilisateurs", None),
    ("Nombre de spectacles", None),
    ("Nombre de contacts", None),
    ("Stockage compris", None),
    ("Emails ou brouillons compris", None),
    ("Crédits William compris", None),
    ("Autres limites", None),
])
add_question(document, "Quel niveau de support est promis pendant la bêta ?", "Support par email, sans délai garanti", lines=2)
add_question(document, "Souhaitez-vous garantir un taux de disponibilité précis ?", "Non pendant la bêta, obligation de moyens")
add_question(document, "Pendant combien de temps les données restent-elles récupérables après résiliation ?", "30 jours")
add_question(document, "Quel plafond de responsabilité souhaitez-vous proposer ?", "Montants HT payés au cours des 12 derniers mois", lines=2)
add_question(document, "ARKENCIEL dispose-t-elle d'une assurance couvrant l'activité logicielle ?", "[ ] Oui    [ ] Non    [ ] À vérifier")

add_page_break(document)
add_heading(document, "4", "Données personnelles et RGPD", "Cette partie précise les responsabilités d'ARKENCIEL, de TaDiff et des compagnies clientes.")
add_question(document, "Qui sera le référent chargé de répondre aux demandes RGPD ?", None)
add_question(document, "Souhaitez-vous désigner un DPO ou un référent interne ?", "[ ] DPO    [ ] Référent interne    [ ] À décider")
add_question(document, "Qui doit être contacté en urgence en cas de fuite ou d'incident de sécurité ?", None)
add_question(document, "Validez-vous la répartition suivante ?", "ARKENCIEL/TaDiff gère les comptes et la facturation. Chaque compagnie reste responsable de ses contacts. TaDiff héberge ces données pour son compte.", lines=2)
add_question(document, "Les données des compagnies peuvent-elles être utilisées pour améliorer commercialement le produit ?", "Non sans information claire et base juridique adaptée", lines=2)
add_question(document, "Les documents et messages des compagnies peuvent-ils servir à entraîner un modèle d'IA ?", "Non par défaut", lines=2)
add_question(document, "Quel délai de suppression appliquer après la fermeture d'un compte ?", "Export pendant 30 jours, puis suppression selon le cycle technique à confirmer")
add_question(document, "Quels journaux de sécurité devons-nous conserver et combien de temps ?", "Connexions et pages authentifiées : 90 jours")
add_question(document, "Qui décide si un incident doit être déclaré à la CNIL et aux personnes concernées ?", None)
add_question(document, "Souhaitez-vous envoyer des communications commerciales aux inscrits ?", "[ ] Oui, avec consentement ou base adaptée et désinscription    [ ] Non")

add_page_break(document)
add_heading(document, "5", "Prestataires techniques", "Compléter les informations connues. Les contrats de traitement des données pourront être vérifiés séparément.")
provider_table = document.add_table(rows=1, cols=4)
provider_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Prestataire", "Usage", "Compte au nom de", "Points à confirmer"]
for index, text in enumerate(headers):
    cell = provider_table.rows[0].cells[index]
    cell.text = text
    set_cell_shading(cell, NAVY)
    set_cell_border(cell, NAVY)
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
set_repeat_table_header(provider_table.rows[0])
providers = [
    ("Supabase", "Comptes, base, documents", "À compléter", "Région, DPA, sauvegardes, suppression"),
    ("Vercel", "Hébergement de l'application", "À compléter", "DPA, région, journaux"),
    ("Stripe", "Paiements et abonnements", "À compléter", "TVA, factures, DPA"),
    ("Mistral / autre IA", "William", "À compléter", "Modèle, pays, conservation, entraînement"),
    ("OpenAI, si utilisé", "Recherche sémantique", "À compléter", "Modèle, conservation, DPA"),
    ("Service email", "Emails de service", "Prestataire à choisir", "Pays, DPA, désinscription"),
    ("Service bancaire", "Connexion bancaire future", "Prestataire à choisir", "Agrément, DPA, données accessibles"),
]
for provider in providers:
    cells = provider_table.add_row().cells
    for index, value in enumerate(provider):
        cells[index].text = value
        set_cell_border(cells[index])
        cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        for run in cells[index].paragraphs[0].runs:
            run.font.size = Pt(7.7)
    set_cell_shading(cells[0], PALE)

add_page_break(document)
add_heading(document, "6", "William", "Ces choix permettent de définir ce que l'assistant peut consulter et ce qu'il peut faire.")
add_question(document, "Quel fournisseur IA souhaitons-nous retenir pendant la phase de test ?", "Mistral à confirmer")
add_question(document, "Quelles données William peut-il consulter ?", "Spectacles, actions, dossiers, dates, aides et finances de la compagnie connectée", lines=2)
add_question(document, "William peut-il lire le contenu complet des documents ?", "Seulement après information de la compagnie et lorsque cela est nécessaire", lines=2)
add_question(document, "Les requêtes et réponses peuvent-elles être conservées ?", "Préciser la durée maximale souhaitée")
add_question(document, "William doit-il être activé séparément pour chaque compte ?", "Oui pendant la phase de rodage")
add_question(document, "Une action proposée par William peut-elle être exécutée sans confirmation ?", "Non, validation humaine obligatoire")
add_question(document, "Quel quota mensuel est inclus dans chaque offre ?", "À définir par formule")
add_question(document, "Que se passe-t-il lorsque le quota est atteint ?", "Blocage jusqu'au renouvellement, sans facturation automatique")

add_page_break(document)
add_heading(document, "7", "Passage d'ARKENCIEL à la société TaDiff")
add_question(document, "À quelle date la société TaDiff doit-elle être constituée ?", None)
add_question(document, "Qui sera son représentant légal ?", None)
add_question(document, "Comment les contrats en cours seront-ils transférés ?", "Information préalable des clients et continuité de leurs droits", lines=2)
add_question(document, "Qui reste responsable des factures, dettes et incidents antérieurs au transfert ?", None, lines=2)
add_question(document, "Quels éléments devront être transférés à TaDiff ?", "Marque, domaine, code, contrats, Stripe, Supabase, Vercel et adresses email", lines=2)
add_question(document, "Les clients devront-ils accepter une nouvelle version des contrats ?", "Oui si les changements sont substantiels")

add_page_break(document)
add_heading(document, "8", "Documents à joindre", "Une copie suffit pour préparer les pages. Les originaux restent conservés par ARKENCIEL.")
checklist = [
    "[ ] Statuts d'ARKENCIEL",
    "[ ] Avis SIRENE ou extrait Kbis",
    "[ ] Récépissé RNA, si association",
    "[ ] Justificatif de l'adresse du siège",
    "[ ] Confirmation du régime de TVA par le comptable",
    "[ ] Attestation d'assurance responsabilité civile professionnelle",
    "[ ] Coordonnées bancaires et identité du titulaire du compte Stripe",
    "[ ] Liste des personnes autorisées à administrer TaDiff",
    "[ ] Contrats ou DPA des principaux prestataires, s'ils sont disponibles",
]
for item in checklist:
    paragraph = document.add_paragraph(item)
    paragraph.paragraph_format.left_indent = Cm(0.4)
    paragraph.paragraph_format.space_after = Pt(6)

add_heading(document, "9", "Remarques ou décisions complémentaires")
remarks = document.add_table(rows=1, cols=1)
remarks.alignment = WD_TABLE_ALIGNMENT.CENTER
remarks_cell = remarks.cell(0, 0)
set_cell_border(remarks_cell)
remarks_cell.text = "À compléter :\n\n\n\n\n\n"

document.add_paragraph()
validation = document.add_table(rows=3, cols=2)
validation.alignment = WD_TABLE_ALIGNMENT.CENTER
for row, label in zip(validation.rows, ["Nom", "Date", "Signature ou validation par email"]):
    row.cells[0].text = label
    row.cells[1].text = "À compléter :"
    set_cell_shading(row.cells[0], PALE)
    for cell in row.cells:
        set_cell_border(cell)

final_note = document.add_paragraph()
final_note.paragraph_format.space_before = Pt(14)
final_note.paragraph_format.space_after = Pt(0)
run = final_note.add_run(
    "Ces réponses serviront à finaliser les documents de TaDiff. Elles devront être relues par un professionnel du droit avant le premier encaissement réel."
)
run.italic = True
run.font.size = Pt(8.5)
run.font.color.rgb = MUTED

document.save(TARGET)
print(TARGET)
